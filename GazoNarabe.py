#!/usr/bin/env python
# -*- coding: utf-8 -*-
# ガゾーナラベ by 片山博文MZ

# sys, os, datetime
import sys, os, datetime
import winreg as reg

class GazoNarabeError(Exception):
    def __init__(self, msg):
        self.msg = msg
    def __str(self):
        return self.msg

# tkinter
import tkinter as tk
import tkinter.ttk as ttk
from tkinter import messagebox
root = tk.Tk()

# Win32 API
import ctypes
from ctypes import c_long, WINFUNCTYPE
from ctypes.wintypes import HWND, UINT, WPARAM, LPARAM, DWORD
WM_DROPFILES = 0x0233
GWL_WNDPROC = -4
DragAcceptFiles = ctypes.windll.shell32.DragAcceptFiles
DragQueryFileW = ctypes.windll.shell32.DragQueryFileW
DragFinish = ctypes.windll.shell32.DragFinish
DragFinish.argtypes = [ ctypes.c_void_p ]
CallWindowProcW = ctypes.windll.user32.CallWindowProcW
CallWindowProcW.argtypes = [ ctypes.c_void_p, HWND, UINT, WPARAM, LPARAM ]
SetWindowLongW = ctypes.windll.user32.SetWindowLongW
GetModuleFileNameW = ctypes.windll.kernel32.GetModuleFileNameW

# window procedure
org_proc = None
dropped = []

@WINFUNCTYPE(c_long, HWND, UINT, WPARAM, LPARAM)
def win_proc(hwnd, msg, wp, lp):
    if msg == WM_DROPFILES:
        nf = DragQueryFileW(wp, -1, None, 0)
        for i in range(nf):
            buf = ctypes.create_unicode_buffer(260)
            if DragQueryFileW(wp, i, buf, 260):
                path = buf.value
                if os.path.isfile(path) and not os.path.isdir(path):
                    dropped.append(path)
        DragFinish(wp)
    return CallWindowProcW(org_proc, hwnd, msg, wp, lp)

# 定数。
NOSPEC = "(指定なし)"
COMPANY_KEY = "SOFTWARE\\Katayama Hirofumi MZ"
SOFT_KEY = COMPANY_KEY + "\\GazoNarabe"

# 空文字列ならNOSPECを返す。文字列を整える。
def NOSPEC_if_empty(text):
    text = text.strip()
    text.translate(str.maketrans({chr(0x0021 + i): chr(0xFF01 + i) for i in range(94)}))
    if text == "":
        return NOSPEC
    return text

# 切り詰める。
def truncate(string, length, ellipsis='...'):
    ret = ""
    i = 0
    for ch in string:
        if i + len(ellipsis) >= length:
            ret += ellipsis
            break;
        ret += ch
        if ord(ch) > 0x7F:
            i += 2
        else:
            i += 1
    return ret

# タイトルを変換する。
def convert_title(title, filename, image_index, the_time, char_limit=0):
    # %F、%N, %nの変換。
    filename = os.path.basename(filename)
    filename.replace("%", "%%")
    filename, ext = os.path.splitext(filename)
    title = title.replace("%F", filename)
    title = title.replace("%N", str(image_index + 1))
    title = title.replace("%n", str(image_index))
    # 日時書式を変換。
    title = the_time.strftime(title)
    import re
    title = re.sub(r'[\\/:*?"<>|]+', '_', title)
    # 長い場合は省略。
    if char_limit > 0:
        title = truncate(title, char_limit - 4, "[...]") + ext
    else:
        title += ext
    return title

def try_int(value, field, another_value = "\x7F"):
    value = str(value).strip()
    value.translate(str.maketrans({chr(0x0021 + i): chr(0xFF01 + i) for i in range(94)}))
    if value == NOSPEC or value == str(another_value):
        return
    try:
        n = int(value)
    except:
        messagebox.showerror("ERROR", "「" + field + "」の欄が間違っています。")
        raise

current_filename = ""

class UISample(ttk.Frame):
    def reset_settings(self):
        self.muki_list = [NOSPEC, "縦向き", "横向き"]
        self.muki_default = NOSPEC
        self.gyousuu_list = ["1", "2", "3", "4"]
        self.gyousuu_default = "2"
        self.retsusuu_list = ["1", "2", "3", "4"]
        self.retsusuu_default = "2"
        self.page_title_list = [NOSPEC, "写真一覧(%N)", "私のアルバム(%N)", "実験結果(%N)"]
        self.page_title_default = "写真一覧(%N)"
        self.page_title_align_list = ["左揃え", "中央揃え", "右揃え"]
        self.page_title_align_default = "中央揃え"
        self.page_size_list = [NOSPEC, "A4", "A3", "A2", "B5", "B4", "B3"]
        self.page_size_default = NOSPEC
        self.image_title_list = [NOSPEC, "%F", "(%N) %F", "%Y年%m月%d日", "%Y年%m月%d日 %H時%M分", "%Y.%m.%d", "%Y.%m.%d %H:%M"]
        self.image_title_default = "(%N) %F"
        self.image_width_list = [NOSPEC, "30", "50", "80", "max"]
        self.image_width_default = "max"
        self.image_height_list = [NOSPEC, "30", "50", "80", "max"]
        self.image_height_default = "max"
        self.cell_height_list = [NOSPEC, "30", "50", "80", "max"]
        self.cell_height_default = "max"
        self.font_name_list = [NOSPEC, "ＭＳ ゴシック", "ＭＳ Ｐゴシック", "ＭＳ 明朝", "ＭＳ Ｐ明朝", "メイリオ", "MS UI Gothic"]
        self.font_name_default = NOSPEC
        self.font_size_list = [NOSPEC, "4", "5", "6", "7"]
        self.font_size_default = NOSPEC
        self.output_name_list = ["写真一覧-%Y.%m.%d", "私のアルバム-%Y.%m.%d", "実験結果-%Y.%m.%d"]
        self.output_name_default = "写真一覧-%Y.%m.%d"
        self.datetime_type_list = ["画像作成日時", "画像更新日時", "docx生成日時"]
        self.datetime_type_default = "画像作成日時"
    def __init__(self, root):
        super().__init__(root, width='620', height='460')
        self.image_ext_list = [".jpg", ".jpeg", ".jpe", ".jfif", ".png", ".gif", ".tif", ".tiff",
                               ".bmp", ".dib"]
        # リストと規定値。
        self.reset_settings()
        # レジストリから設定を読み込む。
        self.first_run = not self.load_settings()
        # フィルターを作成。
        self.filter = "*" + ";*".join(self.image_ext_list)
        # ウィジェットをすべて作成。
        self.createWidgets()
        self.pack()
        # ウィンドウハンドルを取得。
        self.hwnd = self.winfo_id()
        # ドラッグ＆ドロップの準備。
        self.dnd_setup()
        self.dnd_interval = 600
        # 起動コマンドライン引数を処理。
        args = sys.argv[1:]
        self.dnd_notify(args)
        # 初回起動の場合は説明を表示。
        if (self.first_run):
            messagebox.showinfo("ガゾーナラベ",
                "「ガゾーナラベ」は、複数の画像を" +
                "並べて Word 文書ファイル（docx）にするソフトです。\n\n" +
                "基本操作は、画像ファイルをリストに追加して「docx生成」ボタンを" +
                "押すだけです。完成したdocxファイルはデスクトップに作成されます。")
    # ドラッグ＆ドロップされた。
    def dnd_notify(self, filenames):
        for filename in filenames:
            self.insert(filename)
    # ドラッグ＆ドロップを検査。
    def drop_check(self):
        global dropped
        if dropped:
            filenames = dropped
            dropped = []
            self.dnd_notify(filenames)
        self.after(self.dnd_interval, self.drop_check)
    # ドラッグ＆ドロップの準備。
    def dnd_setup(self):
        DragAcceptFiles(self.hwnd, True)
        global org_proc
        org_proc = SetWindowLongW(self.hwnd, GWL_WNDPROC, win_proc)
        self.after_idle(self.drop_check)
    # リストボックスの選択が変わった。
    def listbox_on_sel_change(self, evt=None):
        selection = self.listbox_01.curselection()
        if len(selection) <= 0:
            self.button_03.config(state="disabled")
            self.label_18.image = None
            self.label_18["image"] = None
            return
        self.button_03.config(state="normal")
        filename = self.listbox_01.get(selection[0])
        from PIL import Image, ImageTk
        img = Image.open(filename)
        img = img.resize((48, 48))
        img = ImageTk.PhotoImage(img);
        self.label_18.image = img
        self.label_18["image"] = img
    # ウィジェットをすべて作成。
    def createWidgets(self):
        self.label_01 = ttk.Label(self, text="用紙の向き:", width="", state="normal", )
        self.label_01.place(x=20, y=20)
        self.muki = tk.StringVar()
        self.combobox_01 = ttk.Combobox(self, height="10", state="readonly", width="25", values=self.muki_list, textvariable=self.muki)
        self.combobox_01.place(x=120, y=20)
        self.combobox_01.set(self.muki_default)
        self.label_02 = ttk.Label(self, text="１ページの行数:", width="", state="normal", )
        self.label_02.place(x=20, y=50)
        self.gyousuu = tk.StringVar()
        self.combobox_02 = ttk.Combobox(self, height="10", state="normal", width="25", values=self.gyousuu_list, textvariable=self.gyousuu)
        self.combobox_02.place(x=120, y=50)
        self.combobox_02.set(self.gyousuu_default)
        self.label_03 = ttk.Label(self, text="１ページの列数:", width="", state="normal", )
        self.label_03.place(x=20, y=80)
        self.retsusuu = tk.StringVar()
        self.combobox_03 = ttk.Combobox(self, height="10", state="normal", width="25", values=self.retsusuu_list, textvariable=self.retsusuu)
        self.combobox_03.place(x=120, y=80)
        self.combobox_03.set(self.retsusuu_default)
        self.label_04 = ttk.Label(self, text="ページ見出し:", width="", state="normal", )
        self.label_04.place(x=20, y=110)
        self.page_title = tk.StringVar()
        self.combobox_04 = ttk.Combobox(self, height="10", state="normal", width="25", values=self.page_title_list, textvariable=self.page_title)
        self.combobox_04.place(x=120, y=110)
        self.combobox_04.set(self.page_title_default)
        self.label_05 = ttk.Label(self, text="ページ見出し揃え:", width="", state="normal", )
        self.label_05.place(x=20, y=140)
        self.page_title_align = tk.StringVar()
        self.combobox_05 = ttk.Combobox(self, height="10", state="readonly", width="25", values=self.page_title_align_list, textvariable=self.page_title_align)
        self.combobox_05.place(x=120, y=140)
        self.combobox_05.set(self.page_title_align_default)
        self.label_06 = ttk.Label(self, text="ページサイズ:", width="", state="normal", )
        self.label_06.place(x=20, y=170)
        self.page_size = tk.StringVar()
        self.combobox_06 = ttk.Combobox(self, height="10", state="readonly", width="25", values=self.page_size_list, textvariable=self.page_size)
        self.combobox_06.place(x=120, y=170)
        self.combobox_06.set(self.page_size_default)
        self.label_07 = ttk.Label(self, text="画像タイトル:", width="", state="normal", )
        self.label_07.place(x=20, y=200)
        self.image_title = tk.StringVar()
        self.combobox_07 = ttk.Combobox(self, height="10", state="normal", width="25", values=self.image_title_list, textvariable=self.image_title)
        self.combobox_07.place(x=120, y=200)
        self.combobox_07.set(self.image_title_default)
        self.label_08 = ttk.Label(self, text="画像の幅(mm):", width="", state="normal", )
        self.label_08.place(x=320, y=20)
        self.image_width = tk.StringVar()
        self.combobox_08 = ttk.Combobox(self, height="10", state="normal", width="25", values=self.image_width_list, textvariable=self.image_width)
        self.combobox_08.place(x=420, y=20)
        self.combobox_08.set(self.image_width_default)
        self.label_09 = ttk.Label(self, text="画像の高さ(mm):", width="", state="normal", )
        self.label_09.place(x=320, y=50)
        self.image_height = tk.StringVar()
        self.combobox_09 = ttk.Combobox(self, height="10", state="normal", width="25", values=self.image_height_list, textvariable=self.image_height)
        self.combobox_09.place(x=420, y=50)
        self.combobox_09.set(self.image_height_default)
        self.label_10 = ttk.Label(self, text="セルの高さ(mm):", width="", state="normal", )
        self.label_10.place(x=320, y=80)
        self.cell_height = tk.StringVar()
        self.combobox_10 = ttk.Combobox(self, height="10", state="normal", width="25", values=self.cell_height_list, textvariable=self.cell_height)
        self.combobox_10.place(x=420, y=80)
        self.combobox_10.set(self.cell_height_default)
        self.label_11 = ttk.Label(self, text="フォント名:", width="", state="normal", )
        self.label_11.place(x=320, y=110)
        self.font_name = tk.StringVar()
        self.combobox_11 = ttk.Combobox(self, height="10", state="normal", width="25", values=self.font_name_list, textvariable=self.font_name)
        self.combobox_11.place(x=420, y=110)
        self.combobox_11.set(self.font_name_default)
        self.label_12 = ttk.Label(self, text="フォントサイズ(mm):", width="", state="normal", )
        self.label_12.place(x=320, y=140)
        self.font_size = tk.StringVar()
        self.combobox_12 = ttk.Combobox(self, height="10", state="normal", width="25", values=self.font_size_list, textvariable=self.font_size)
        self.combobox_12.place(x=420, y=140)
        self.combobox_12.set(self.font_size_default)
        self.label_17 = ttk.Label(self, text="日付の種類:", width="", state="normal", )
        self.label_17.place(x=320, y=170)
        self.datetime_type = tk.StringVar()
        self.combobox_14 = ttk.Combobox(self, height="10", state="readonly", width="25", values=self.datetime_type_list, textvariable=self.datetime_type)
        self.combobox_14.place(x=420, y=170)
        self.combobox_14.set(self.datetime_type_default)
        self.label_13 = ttk.Label(self, text="出力ファイル名:", width="", state="normal", )
        self.label_13.place(x=320, y=200)
        self.output_name = tk.StringVar()
        self.combobox_13 = ttk.Combobox(self, height="10", state="normal", width="25", values=self.output_name_list, textvariable=self.output_name)
        self.combobox_13.place(x=420, y=200)
        self.combobox_13.set(self.output_name_default)
        self.label_14 = ttk.Label(self, text="%F:ファイル名, %N:1から連番, %n:0から連番, %Y:西暦年, %m:月, %d:日, %H:時, %M:分, %%:%", width="200", state="normal", )
        self.label_14.place(x=20, y=230)
        self.label_15 = ttk.Label(self, text="docx生成の際には、ひな形としてtemplate.docxの設定が使用されます。", width="200", state="normal", )
        self.label_15.place(x=20, y=245)
        self.label_16 = ttk.Label(self, text="並べる画像ファイルのリスト: (ファイルドロップが可能です)", width="200", state="normal", )
        self.label_16.place(x=20, y=280)
        self.total_number = tk.StringVar()
        self.label_17 = ttk.Label(self, text="　　　　　　", width="200", state="normal", textvariable=self.total_number)
        self.label_17.place(x=320, y=280)
        self.label_18 = ttk.Label(self, width="32", state="normal", image="")
        self.label_18.place(x=400, y=245)
        self.group1 = tk.Frame(self)
        self.group1.place(x=20, y=300)
        self.listbox_01 = tk.Listbox(self.group1, width=92, height=7, selectmode=tk.EXTENDED, activestyle='none')
        self.listbox_01.pack(side=tk.LEFT, fill=tk.Y)
        self.listbox_01.config(borderwidth=2)
        self.vscrollbar = tk.Scrollbar(self.group1, orient=tk.VERTICAL, command=self.listbox_01.yview)
        self.vscrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        self.listbox_01.config(yscrollcommand = self.vscrollbar.set)
        self.listbox_01.bind('<<ListboxSelect>>', self.listbox_on_sel_change)
        self.button_01 = ttk.Button(self, command = self.commandAddFiles, text="リストに追加...", width="15", state="normal", )
        self.button_01.place(x=20, y=420)
        self.button_02 = ttk.Button(self, command = self.commandMoveUp, text="↑", width="8", state="normal", )
        self.button_02.place(x=140, y=420)
        self.button_03 = ttk.Button(self, command = self.commandMoveDown, text="↓", width="8", state="normal", )
        self.button_03.place(x=200, y=420)
        self.button_03 = ttk.Button(self, command = self.commandDeleteItems, text="選択を削除", width="15", state="normal", )
        self.button_03.place(x=280, y=420)
        self.button_04 = ttk.Button(self, command = self.commandOK, text="docx生成", width="", state="normal", )
        self.button_04.place(x=410, y=420)
        self.button_05 = ttk.Button(self, command = self.commandExit, text="終了", width="", state="normal", )
        self.button_05.place(x=520, y=420)
        self.button_06 = ttk.Button(self, command = self.commandResetSettings, text="設定の初期化", width="", state="normal", )
        self.button_06.place(x=510, y=270)
        self.update_count()
    # 個数を更新。
    def update_count(self):
        if self.listbox_01.size() <= 0:
            self.button_03.config(state="disabled")
            self.total_number.set("")
            self.listbox_on_sel_change()
            return
        self.total_number.set("全部で" + str(self.listbox_01.size()) + "個")
        self.button_03.config(state="normal")
        self.listbox_on_sel_change()
    # 挿入。
    def insert(self, filename):
        ext = os.path.splitext(filename)[1].lower()
        if ext in self.image_ext_list:
            self.listbox_01.insert(tk.END, filename)
        self.update_count()
    # 「追加」ボタンを押した。
    def commandAddFiles(self):
        from tkinter import filedialog
        names = filedialog.askopenfilenames(initialdir=".", title="追加する画像ファイル", \
            filetypes = (("画像ファイル", self.filter), ("すべてのファイル", "*.*")))
        filenames = root.tk.splitlist(names)
        for filename in filenames:
            self.insert(filename)
    # 「選択を削除」ボタンを押した。
    def commandDeleteItems(self):
        items = self.listbox_01.curselection()
        for item in reversed(list(items)):
            self.listbox_01.delete(item)
        self.update_count();
    # 「↑」ボタンを押した。
    def commandMoveUp(self):
        items = list(self.listbox_01.curselection())
        items.sort()
        if items[0] == 0:
            return;
        for item in items:
            text0 = self.listbox_01.get(item - 1)
            text1 = self.listbox_01.get(item)
            self.listbox_01.delete(item - 1)
            self.listbox_01.delete(item - 1)
            self.listbox_01.insert(item - 1, text0)
            self.listbox_01.insert(item - 1, text1)
        for item in items:
            self.listbox_01.selection_clear(item)
        for item in items:
            self.listbox_01.selection_set(item - 1)
    # 「↓」ボタンを押した。
    def commandMoveDown(self):
        items = list(self.listbox_01.curselection())
        items.sort()
        items.reverse()
        if items[0] == self.listbox_01.size() - 1:
            return;
        for item in items:
            text0 = self.listbox_01.get(item)
            text1 = self.listbox_01.get(item + 1)
            self.listbox_01.delete(item)
            self.listbox_01.delete(item)
            self.listbox_01.insert(item, text0)
            self.listbox_01.insert(item, text1)
        for item in items:
            self.listbox_01.selection_clear(item)
        for item in items:
            self.listbox_01.selection_set(item + 1)
    # リストボックスからファイルのリストを取得する。
    def get_file_list(self):
        file_list = []
        size = self.listbox_01.size()
        for i in range(size):
            text = self.listbox_01.get(i)
            file_list.append(text)
        return file_list
    # docxファイルを生成する。
    def generate_docx(self):
        import docx
        from docx.table import Table
        from docx.enum.section import WD_ORIENT
        from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.style import WD_STYLE_TYPE
        from docx.shared import Mm, Inches, Pt
        from docx.oxml.ns import qn

        # リストボックスからファイルのリストを取得する。
        file_list = self.get_file_list()

        # 実行モジュールのパスを取得。
        buf = ctypes.create_unicode_buffer(260)
        GetModuleFileNameW(None, buf, 260)
        mod_path = buf.value

        # 最初の段落を取得する。
        dir = os.path.dirname(os.getcwd() + "/" + __file__)
        file = dir + "/template.docx"
        if not os.path.isfile(file):
            dir = os.path.dirname(os.getcwd() + "/" + __file__)
            file = dir + "/../template.docx"
        elif not os.path.isfile(file):
            dir = os.path.dirname(mod_path)
            file = dir + "/../template.docx"
        try:
            document = docx.Document(file)
        except:
            messagebox.showerror("ERROR",
                "生成に必要なファイル「template.docx」が見つかりません。")
            return False

        para = document.paragraphs[0]

        #for style in document.styles:
        #    print(style)

        # フォント名。
        the_font_name = self.font_name_default
        if the_font_name == NOSPEC:
            the_font_name = para.style.font.name        # 指定なしなら最初の段落を使う。

        # フォントサイズ(mm)。
        the_font_size = self.font_size_default
        if the_font_size == NOSPEC:
            # 指定なしなら最初の段落を使う。
            if para.style.font.size == None:
                the_font_size = int(document.styles['Normal'].font.size.mm)
            else:
                the_font_size = int(para.style.font.size.mm)
        else:
            the_font_size = int(the_font_size)
        #print("the_font_size: " + str(the_font_size))

        # 文書をクリア。
        document._body.clear_content()

        # すべてのセクションにて。
        section = document.sections[-1]

        # ページサイズの指定を取得する(mm)。
        PAGE_SIZE_INFO = [
            ["A2", 420, 594],
            ["A3", 297, 420],
            ["A4", 210, 297],
            ["A5", 148, 210],
            ["B2", 515, 728],
            ["B3", 364, 515],
            ["B4", 257, 364],
            ["B5", 182, 257]]
        for info in PAGE_SIZE_INFO:
            if info[0] == self.page_size_default:
                section.page_width = Mm(info[1])
                section.page_height = Mm(info[2])
                break

        # 横向き指定なら横向きにする。
        if self.muki_default == "横向き":
            new_width, new_height = section.page_height, section.page_width
            section.orientation = WD_ORIENT.LANDSCAPE
            section.page_height = new_height
            section.page_width = new_width

        # 実際のページサイズ(mm)。
        the_page_height = section.page_height.mm
        the_page_width = section.page_width.mm

        # 余白を計算する(mm)。
        x_margin = section.left_margin.mm + section.right_margin.mm
        y_margin = section.top_margin.mm + section.bottom_margin.mm

        # 行数と列数。
        lines = int(self.gyousuu_default)
        columns = int(self.retsusuu_default)

        # 表の印刷可能領域(mm)。
        table_width = the_page_width - x_margin
        table_height = the_page_height - y_margin
        table_height -= the_font_size # ページの最後に１行残す。
        if self.page_title_default != NOSPEC:
            table_height -= the_font_size

        # セルの最大サイズ(mm)。
        cell_max_width = int(table_width / columns)
        cell_max_height = int(table_height / lines)
        #print(cell_max_width)

        # セルの中身の最大サイズ(mm)。
        contents_width = table_width
        contents_height = table_height
        if self.image_title_default != NOSPEC:
            contents_height -= the_font_size * 2 * lines
        contents_width = int(contents_width / columns)
        contents_height = int(contents_height / lines)

        # 文字数制限。
        char_limit = int(contents_width * 1.7 / the_font_size)
        #print("char_limit: " + str(char_limit))

        # ページセル数。
        page_cells = lines * columns

        if len(file_list) <= 0:
            raise GazoNarabeError("画像ファイルリストが空です。")

        # 実際に生成する。
        image_index = 0
        col_index = 0
        row_index = 0
        page_number = 0
        table = None

        for i in range(0, len(file_list)):
            filename = file_list[image_index]
            # 日時文字列の処理。
            if self.datetime_type_default == "画像作成日時":
                try:
                    from PIL import Image
                    s = Image.open(filename)._getexif()[36867]
                    s = s.replace(":", "-", 2)
                    the_time = datetime.datetime.fromisoformat(s)
                    #print("EXIF: " + filename + " | " + str(the_time))
                except:
                    timestamp = os.path.getctime(filename)
                    the_time = datetime.datetime.fromtimestamp(timestamp)
                    #print("non-EXIF: " + filename + " | " + str(the_time))
            elif self.datetime_type_default == "画像更新日時":
                timestamp = os.path.getmtime(filename)
                the_time = datetime.datetime.fromtimestamp(timestamp)
            elif self.datetime_type_default == "docx生成日時":
                the_time = datetime.datetime.now()
            else:
                the_time = datetime.datetime.now()
            if i % page_cells == 0:
                # ページの最初のセル。
                if i != 0:
                    # 文書の最初のセルでなければ、直前に作成した行の高さを調整する。
                    if self.cell_height_default != NOSPEC:
                        if self.cell_height_default == "max":
                            for row in table.rows:
                                row.height = Mm(int(cell_max_height))
                        else:
                            for row in table.rows:
                                row.height = Mm(int(self.cell_height_default))
                    # ページ区切りを挿入。
                    document.add_page_break()
                # 必要ならば、ページ見出しを追加する。
                if self.page_title_default != NOSPEC:
                    # 段落を挿入する。
                    para = document.add_paragraph()
                    title = convert_title(self.page_title_default, "", page_number, the_time)
                    run = para.add_run(title)
                    # フォント設定。
                    run.font.name = the_font_name
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), the_font_name)
                    run.font.size = Mm(the_font_size)
                    the_align = self.page_title_align_default
                    if the_align == "左揃え":
                        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    elif the_align == "中央揃え":
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    elif the_align == "右揃え":
                        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                # 表を挿入する。行はゼロ個。
                table = document.add_table(rows=0, cols=columns)
                table.allow_autofit = False
                # 表を中央揃えにする。
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
            # 行の直前なら
            if i % columns == 0:
                # 行を追加する。
                row_cells = table.add_row().cells
            # セルを取得し、縦に中央揃えする。
            cell = row_cells[col_index]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            # セルの最初の段落を左右中央揃えにする。
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # 画像を貼り付ける。
            run = para.add_run()
            # 画像サイズを得る。
            requested_width = self.image_width_default
            requested_height = self.image_height_default
            real_image_width = real_image_height = 0
            from PIL import Image
            with Image.open(filename) as image:
                real_image_width = image.width
                real_image_height = image.height
            if real_image_width * real_image_height != 0:
                if requested_width == "max" or requested_height == "max":
                    aspect0 = float(contents_height) / float(contents_width)
                    aspect1 = float(real_image_height) / float(real_image_width)
                    if aspect0 < aspect1:
                        requested_height = contents_height
                        requested_width = requested_height / aspect1
                    else:
                        requested_width = contents_width
                        requested_height = requested_width * aspect1
                    requested_width *= 0.98
                    requested_height *= 0.98
            global current_filename
            current_filename = filename
            if requested_width == NOSPEC and requested_height == NOSPEC:
                run.add_picture(filename)
            elif requested_width != NOSPEC and requested_height == NOSPEC:
                run.add_picture(filename, width = Mm(int(requested_width)))
            elif requested_width == NOSPEC and requested_height != NOSPEC:
                run.add_picture(filename, height = Mm(int(requested_height)))
            elif requested_width != NOSPEC and requested_height != NOSPEC:
                run.add_picture(filename, width = Mm(int(requested_width)), height = Mm(int(requested_height)))
            else:
                raise
            # 画像のタイトルを取得する。
            image_title = self.image_title_default
            if image_title != NOSPEC:
                image_title = convert_title(image_title, filename, image_index, the_time, char_limit)
                # タイトル用の段落を追加する。
                para = cell.add_paragraph()
                # タイトルを入れる。
                run = para.add_run(image_title + "\n")
                # タイトルのフォント設定。
                run.font.name = the_font_name
                run._element.rPr.rFonts.set(qn('w:eastAsia'), the_font_name)
                run.font.size = Mm(the_font_size)
                # 中央揃えにする。
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # インデックスを更新する。
            image_index += 1
            col_index += 1
            col_index %= columns
            if col_index == 0:
                row_index += 1
                if (image_index % page_cells == 0):
                    page_number += 1
        # もし、セルの高さが未指定でなければ
        if self.cell_height_default != NOSPEC:
            if self.cell_height_default == "max":
                for row in table.rows:
                    row.height = Mm(int(cell_max_height))
            else:
                for row in table.rows:
                    row.height = Mm(int(self.cell_height_default))
        # 変更内容を保存する。
        title = convert_title(self.output_name_default, "", 0, the_time) + ".docx"
        dir = os.path.join(os.path.join(os.environ['USERPROFILE']), 'Desktop')
        filename = dir + "/" + title
        try:
            document.save(filename)
        except:
            raise GazoNarabeError("ファイル「" + filename + "」はロックされていたため、保存に失敗しました。")
        messagebox.showinfo("ガゾーナラベ",
            "デスクトップにファイル「" + title + "」の保存に成功しました。")
        return True
    # 「生成する」ボタンを押した。
    def commandOK(self):
        # 設定を取り出す。
        self.muki_default = self.muki.get()
        self.gyousuu_default = self.gyousuu.get()
        self.retsusuu_default = self.retsusuu.get()
        self.page_title_default = self.page_title.get()
        self.page_title_align_default = self.page_title_align.get()
        self.page_size_default = self.page_size.get()
        self.image_title_default = self.image_title.get()
        self.image_width_default = self.image_width.get()
        self.image_height_default = self.image_height.get()
        self.cell_height_default = self.cell_height.get()
        self.font_name_default = self.font_name.get()
        self.font_size_default = self.font_size.get()
        self.output_name_default = self.output_name.get()
        self.datetime_type_default = self.datetime_type.get()

        try_int(self.gyousuu_default, "行数")
        try_int(self.retsusuu_default, "列数")
        try_int(self.image_width_default, "画像の幅", "max")
        try_int(self.image_height_default, "画像の高さ", "max")
        try_int(self.cell_height_default, "セルの高さ", "max")
        try_int(self.font_size_default, "フォントサイズ")

        # 空文字列ならNOSPECにする。
        self.muki.set(NOSPEC_if_empty(self.muki_default))
        self.gyousuu.set(NOSPEC_if_empty(self.gyousuu_default))
        self.retsusuu.set(NOSPEC_if_empty(self.retsusuu_default))
        self.page_title.set(NOSPEC_if_empty(self.page_title_default))
        self.page_title_align.set(NOSPEC_if_empty(self.page_title_align_default))
        self.page_size.set(NOSPEC_if_empty(self.page_size_default))
        self.image_title.set(NOSPEC_if_empty(self.image_title_default))
        self.image_width.set(NOSPEC_if_empty(self.image_width_default))
        self.image_height.set(NOSPEC_if_empty(self.image_height_default))
        self.cell_height.set(NOSPEC_if_empty(self.cell_height_default))
        self.font_name.set(NOSPEC_if_empty(self.font_name_default))
        self.font_size.set(NOSPEC_if_empty(self.font_size_default))
        self.output_name.set(NOSPEC_if_empty(self.output_name_default))
        self.datetime_type.set(NOSPEC_if_empty(self.datetime_type_default))

        # 設定内容を元に生成する。
        try:
            self.generate_docx()
        except Exception as err:
            global current_filename
            if str(err) != "":
                messagebox.showerror("ERROR: ガゾーナラベ", str(err))
            elif current_filename != "":
                messagebox.showerror("ERROR: ガゾーナラベ", "ファイル「" + current_filename + "」の処理に失敗しました。")
            else:
                messagebox.showerror("ERROR: ガゾーナラベ", "何らかの処理に失敗しました。")
        # 設定を保存する。
        self.save_settings()
    # 「終了」ボタンを押した。
    def commandExit(self):
        root.destroy()
    # 「設定の初期化」ボタンを押した。
    def commandResetSettings(self):
        self.reset_settings()
        self.save_settings()
        messagebox.showinfo("ガゾーナラベ", "設定を初期化しました。")
        root.destroy()
    # １つ設定を読み込む。
    def read_settings(self, key, name, the_list, value):
        try:
            count = int(reg.QueryValueEx(key, name + "_count")[0])
            if count >= 0:
                the_list.clear()
                for i in range(count):
                    value = reg.QueryValueEx(key, name + "_" + str(i))[0]
                    the_list.append(value)
            return reg.QueryValueEx(key, name + "_value")[0]
        except:
            return value
    # １つ設定を書き込む。
    def write_settings(self, key, name, the_list, value):
        if value not in the_list:
            the_list.append(value)
        reg.SetValueEx(key, name + "_value", 0, reg.REG_SZ, value)
        reg.SetValueEx(key, name + "_count", 0, reg.REG_SZ, str(len(the_list)))
        for i, item in enumerate(the_list):
            reg.SetValueEx(key, name + "_" + str(i), 0, reg.REG_SZ, item)
        return True
    # レジストリから設定を読み込む。
    def load_settings(self):
        first_run = True
        try:
            with reg.OpenKeyEx(reg.HKEY_CURRENT_USER, SOFT_KEY, 0, reg.KEY_READ|reg.KEY_WOW64_64KEY) as soft_key:
                first_run = False
                self.muki_default = self.read_settings(soft_key, "muki", self.muki_list, self.muki_default)
                self.gyousuu_default = self.read_settings(soft_key, "gyousuu", self.gyousuu_list, self.gyousuu_default)
                self.retsusuu_default = self.read_settings(soft_key, "retsusuu", self.retsusuu_list, self.retsusuu_default)
                self.page_title_default = self.read_settings(soft_key, "page_title", self.page_title_list, self.page_title_default)
                self.page_title_align_default = self.read_settings(soft_key, "page_title_align_list", self.page_title_align_list, self.page_title_align_default)
                self.page_size_default = self.read_settings(soft_key, "page_size", self.page_size_list, self.page_size_default)
                self.image_title_default = self.read_settings(soft_key, "image_title", self.image_title_list, self.image_title_default)
                self.image_width_default = self.read_settings(soft_key, "image_width", self.image_width_list, self.image_width_default)
                self.image_height_default = self.read_settings(soft_key, "image_height", self.image_height_list, self.image_height_default)
                self.cell_height_default = self.read_settings(soft_key, "cell_height", self.cell_height_list, self.cell_height_default)
                self.font_name_default = self.read_settings(soft_key, "font_name", self.font_name_list, self.font_name_default)
                self.font_size_default = self.read_settings(soft_key, "font_size", self.font_size_list, self.font_size_default)
                self.output_name_default = self.read_settings(soft_key, "output_name", self.output_name_list, self.output_name_default)
                self.datetime_type_default = self.read_settings(soft_key, "datetime_type", self.datetime_type_list, self.datetime_type_default)
        except:
            pass
        return not first_run
    # レジストリに設定を書き込む。
    def save_settings(self):
        try:
            with reg.CreateKeyEx(reg.HKEY_CURRENT_USER, COMPANY_KEY, 0, reg.KEY_WRITE|reg.KEY_WOW64_64KEY) as company_key:
                with reg.CreateKeyEx(company_key, "GazoNarabe", 0, reg.KEY_WRITE|reg.KEY_WOW64_64KEY) as soft_key:
                    self.write_settings(soft_key, "muki", self.muki_list, self.muki_default)
                    self.write_settings(soft_key, "gyousuu", self.gyousuu_list, self.gyousuu_default)
                    self.write_settings(soft_key, "retsusuu", self.retsusuu_list, self.retsusuu_default)
                    self.write_settings(soft_key, "page_title", self.page_title_list, self.page_title_default)
                    self.write_settings(soft_key, "page_title_align_list", self.page_title_align_list, self.page_title_align_default)
                    self.write_settings(soft_key, "page_size", self.page_size_list, self.page_size_default)
                    self.write_settings(soft_key, "image_title", self.image_title_list, self.image_title_default)
                    self.write_settings(soft_key, "image_width", self.image_width_list, self.image_width_default)
                    self.write_settings(soft_key, "image_height", self.image_height_list, self.image_height_default)
                    self.write_settings(soft_key, "cell_height", self.cell_height_list, self.cell_height_default)
                    self.write_settings(soft_key, "font_name", self.font_name_list, self.font_name_default)
                    self.write_settings(soft_key, "font_size", self.font_size_list, self.font_size_default)
                    self.write_settings(soft_key, "output_name", self.output_name_list, self.output_name_default)
                    self.write_settings(soft_key, "datetime_type", self.datetime_type_list, self.datetime_type_default)
                    return True
        except:
            return False

# 主処理。
root.title('ガゾーナラベ version 0.8 by 片山博文MZ')
root.geometry("620x460")
root.resizable(width=False, height=False)
frame = UISample(root)
root.mainloop()

with reg.CreateKeyEx(reg.HKEY_CURRENT_USER, COMPANY_KEY, 0, reg.KEY_WRITE|reg.KEY_WOW64_64KEY) as company_key:
    with reg.CreateKeyEx(company_key, "GazoNarabe", 0, reg.KEY_WRITE|reg.KEY_WOW64_64KEY) as soft_key:
        pass
